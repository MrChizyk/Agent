"""
Витягування тексту з PDF та DOCX файлів для подальшого аналізу через AI.
"""
import PyPDF2
import docx

from config import MAX_DOCUMENT_CHARS


def extract_text_from_pdf(file_path: str) -> str:
    text_parts = []
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    text = "\n".join(text_parts).strip()
    return _truncate(text)


def extract_text_from_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Також витягуємо текст із таблиць, якщо вони є в документі
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)

    text = "\n".join(paragraphs).strip()
    return _truncate(text)


def _truncate(text: str) -> str:
    if len(text) > MAX_DOCUMENT_CHARS:
        return text[:MAX_DOCUMENT_CHARS] + "\n\n[...текст обрізано через великий обсяг...]"
    return text
